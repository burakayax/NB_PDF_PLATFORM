import pytesseract
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import PyPDF2
import io
import os
import platform
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
import logging
from typing import List, Dict, Any, Tuple, Optional

logger = logging.getLogger(__name__)

try:
    from pypdf.errors import PdfReadError as _PdfReadError
except ImportError:  # pragma: no cover
    try:
        from PyPDF2.errors import PdfReadError as _PdfReadError
    except ImportError:  # pragma: no cover
        _PdfReadError = Exception  # type: ignore[misc,assignment]

# --- YOLLAR ---
_tess_which = shutil.which("tesseract")
_tess_fallback = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if _tess_which:
    pytesseract.pytesseract.tesseract_cmd = _tess_which
elif os.path.isfile(_tess_fallback):
    pytesseract.pytesseract.tesseract_cmd = _tess_fallback
else:
    raise EnvironmentError(
        "Tesseract bulunamadı. Lütfen Tesseract-OCR'ı kurun ve PATH'e ekleyin: "
        "https://github.com/UB-Mannheim/tesseract/wiki"
    )
base_dir = os.path.dirname(os.path.abspath(__file__))
# Poppler klasörünün proje ana dizinindeki Library/bin içinde olduğu varsayılır
poppler_bin_path = os.path.join(base_dir, "..", "Library", "bin")

# Taranmış form / ekran görüntüsü PDF'ler için sayfa bölümleme (PSM 3: otomatik)
_TESSERACT_CONFIG = r"--oem 3 --psm 3"


def _pdf_open_user_message(where: str, exc: BaseException, basename: str) -> str:
    """Dosya yolu sızdırmadan kısa teknik özet (istemci/tool_errors ile uyumlu)."""
    ename = type(exc).__name__
    msg = (str(exc) or "").strip().replace("\n", " ")
    if len(msg) > 200:
        msg = msg[:197] + "..."
    return f"[{ename}] {where} ({basename}): {msg}"


def _pymupdf_page_count(pdf_path: str, password: Optional[str] = None) -> int:
    """pypdf başarısız olursa sayfa sayısı için PyMuPDF yedek."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        if doc.needs_pass:
            pwd = (password or "").strip()
            if doc.authenticate(pwd):
                pass
            elif doc.authenticate(""):
                pass
            else:
                raise Exception("PyMuPDF: PDF parolası gerekli veya hatalı.")
        return doc.page_count
    finally:
        doc.close()


def _open_pdf_reader(pdf_path: str, password: Optional[str] = None, context: str = "Bu işlem") -> PyPDF2.PdfReader:
    """PDF okuyucu: her zaman rb akışı; boş kullanıcı parolası için decrypt(''); açıklamalı hatalar."""
    basename = os.path.basename(pdf_path)
    reader: PyPDF2.PdfReader
    try:
        with open(pdf_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh, strict=False)
    except _PdfReadError as e:
        raise Exception(_pdf_open_user_message("pypdf.PdfReader", e, basename)) from e
    except OSError as e:
        raise Exception(_pdf_open_user_message("dosya okuma", e, basename)) from e
    except Exception as e:
        raise Exception(_pdf_open_user_message("pypdf.PdfReader", e, basename)) from e

    try:
        if reader.is_encrypted:
            pwd = (password or "").strip()
            decrypted = 0
            if pwd:
                decrypted = reader.decrypt(pwd)
            if not decrypted:
                decrypted = reader.decrypt("")
            if not decrypted:
                if not pwd:
                    raise Exception(
                        f"{context} için seçtiğiniz dosya şifreli: {basename}\n"
                        "Lütfen dosya için şifre girin."
                    )
                raise Exception(f"{context} için girilen şifre hatalı: {basename}")
        return reader
    except Exception as e:
        err_text = str(e).lower()
        if (
            "şifreli" in err_text
            or "girilen şifre hatalı" in err_text
            or "girilen pdf parolası hatalı" in err_text
            or "lütfen dosya için şifre" in err_text
        ):
            raise
        raise Exception(_pdf_open_user_message("PDF şifre çözümü", e, basename)) from e


def _apply_output_pdf_password(output_path: str, output_password: Optional[str]) -> None:
    """Yazılmış PDF çıktısına görüntüleyici parolası uygular (pikepdf ile).
    Parola boşsa dosyaya dokunmaz; şifreleme modülü akışını sadeleştirir.
    pikepdf yoksa veya dosya kilitliyse çağıran kullanıcıya anlamlı hata göstermelidir."""
    if not output_password:
        return
    try:
        import pikepdf
    except ImportError as e:
        raise Exception("PDF çıktı şifreleme için 'pikepdf' gerekli.") from e

    owner = output_password
    out_dir = os.path.dirname(output_path) or None
    fd, temp_output_path = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        with pikepdf.open(output_path) as pdf:
            pdf.save(
                temp_output_path,
                encryption=pikepdf.Encryption(
                    user=output_password,
                    owner=owner,
                    R=6,
                    allow=pikepdf.Permissions(extract=False),
                ),
            )
        os.replace(temp_output_path, output_path)
    finally:
        if os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except Exception:
                pass


def classify_pdf_password_requirement(pdf_path: str) -> Tuple[bool, Dict[str, Any]]:
    """
    (requires_non_empty_password, diagnostics)

    Requires True ise kullanıcı arayüzünde açılış şifresi istenmeli.
    Bazı PDF'ler Encrypt sözlüğü nedeniyle PyPDF'de `is_encrypted=True` görünürken
    boş kullanıcı parolasıyla veya görüntüleyici yüzünden parolasız açılabilir;
    ikinci olarak PyMuPDF ile doğrulanır; PyPDF başarısız olursa yalnızca PyMuPDF denenir.
    diagnostics: parola sızdırmadan kısa teşhis bilgisi (log / inspect API).
    """
    basename = os.path.basename(pdf_path)
    diagnostics: Dict[str, Any] = {"file": basename, "engines": {}}

    def _finalize(requires_pw: bool, reason: str) -> Tuple[bool, Dict[str, Any]]:
        diagnostics["requires_password"] = requires_pw
        diagnostics["classification_reason"] = reason
        log_fn = logger.info if requires_pw else logger.debug
        log_fn("[pdf-password] %s requires_password=%s %s", basename, requires_pw, diagnostics)
        return requires_pw, diagnostics

    reader: Optional[PyPDF2.PdfReader] = None
    try:
        with open(pdf_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh, strict=False)
    except _PdfReadError as e:
        diagnostics["engines"]["pypdf"] = {
            "open_ok": False,
            "detail": _pdf_open_user_message("classify_pdf pypdf", e, basename),
        }
        return _finalize(*_fallback_classify_via_pymupdf(pdf_path, diagnostics, pdf_read_error=e))
    except Exception as e:
        diagnostics["engines"]["pypdf"] = {
            "open_ok": False,
            "detail": _pdf_open_user_message("classify_pdf pypdf", e, basename),
        }
        return _finalize(*_fallback_classify_via_pymupdf(pdf_path, diagnostics, pdf_read_error=e))

    diagnostics["engines"]["pypdf"] = {
        "is_encrypted_flag": bool(reader.is_encrypted),
    }

    if not reader.is_encrypted:
        diagnostics["engines"]["pypdf"]["note"] = "Encrypt sözlüğü / is_encrypted yok → parola gerekmiyor."
        return _finalize(False, "pypdf_not_encrypted")

    rc_empty = 0
    try:
        rc_empty = reader.decrypt("")
    except Exception as dex:
        diagnostics["engines"]["pypdf"]["decrypt_empty_error"] = f"{type(dex).__name__}:{dex}"
    diagnostics["engines"]["pypdf"]["decrypt_empty_rc"] = rc_empty

    if rc_empty not in (0, False, None):
        diagnostics["engines"]["pypdf"]["note"] = f"Boş kullanıcı parolasıyla decrypt işlemi döndü: {rc_empty}"
        return _finalize(False, "pypdf_decrypt_empty_succeeded")

    lazy_ok, lazy_msg = _pypdf_attempt_read_plain_content(reader)
    diagnostics["engines"]["pypdf"]["lazy_read"] = lazy_msg
    if lazy_ok:
        diagnostics["engines"]["pypdf"][
            "note"
        ] = "Decrypt boş ile 0 görünse bile içerik okunabildi (şifresiz kullanıcı parolası / metadata)."
        return _finalize(False, "pypdf_lazy_content_readable")

    pym_needs_pw, pym_meta = pymupdf_requires_non_empty_password(pdf_path)
    diagnostics["engines"]["pymupdf"] = pym_meta

    if not pym_needs_pw:
        diagnostics[
            "note"
        ] = "PyPDF parola gerektiriyor görünürken PyMuPDF açılış için boş kullanıcı parolası yeterli (veya parola gerekmiyor)."
        return _finalize(False, "pymupdf_relaxed_false_positive_guard")

    diagnostics["engines"]["note"] = "Her iki katmanda da kullanıcı açılış şifresi gerekiyor kabulü."
    return _finalize(True, "password_required")


def pymupdf_requires_non_empty_password(pdf_path: str) -> Tuple[bool, Dict[str, Any]]:
    """
    True ise boş şifreyle açılamıyor, dolayısı kullanıcıdan parola beklenmeli.
    Fitiz (PyMuPDF) yüklenmezse güvenli tarafta kalınır ve (True, {error:...}) döner.
    """
    meta: Dict[str, Any] = {}
    try:
        import fitz  # PyMuPDF
    except ImportError:
        meta["error"] = "PyMuPDF (fitz) not installed — cannot classify"
        logger.warning("[pdf-password] pymupdf import failed for %s", os.path.basename(pdf_path))
        return True, meta

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        meta["open_error"] = f"{type(e).__name__}:{e}"
        return True, meta

    try:
        meta["needs_pass"] = bool(doc.needs_pass)
        if not doc.needs_pass:
            meta["note"] = "needs_pass=false — açılışta parola istemi yok."
            return False, meta

        auth_attempts = []
        for candidate in ("", b""):
            try:
                ok = bool(doc.authenticate(candidate))
                auth_attempts.append({"candidate_type": type(candidate).__name__, "ok": ok})
                if ok:
                    meta["authenticate"] = "empty_accepted"
                    meta["attempts"] = auth_attempts
                    meta[
                        "note"
                    ] = "needs_pass olsa bile boş kullanıcı parolası kabul edildi (şifreyi gerektiren arayüz gereksiz)."
                    return False, meta
            except Exception as ae:
                auth_attempts.append(
                    {"candidate_type": type(candidate).__name__, "error": f"{type(ae).__name__}:{ae}"}
                )
        meta["attempts"] = auth_attempts

        probe_ok, probe_msg = _pymupdf_probe_first_page_text(doc)
        meta["lazy_page_probe"] = probe_msg
        if probe_ok:
            meta[
                "note"
            ] = "authenticate boş ile başarısız görünse bile ilk sayfa okunuyor → parola zorunluluğu yok varsayılıyor."
            return False, meta

        meta[
            "note"
        ] = "needs_pass ve boş parola reddi — kullanıcı açılış parolası gerekli kabulü."
        return True, meta
    finally:
        doc.close()


def _pymupdf_probe_first_page_text(doc: Any) -> Tuple[bool, str]:
    """İlk sayfadaki metnin okunup okunmadığını dener — taranmış görüntü PDF'lerinde metin olmayabilir (False)."""
    try:
        n = getattr(doc, "page_count", 0) or 0
        if n == 0:
            return False, "page_count_zero"
        txt = doc.load_page(0).get_text() or ""
        stripped = (txt or "").strip()
        return (len(stripped) > 0), f"non_empty_chars={len(stripped)}"
    except Exception as exc:
        return False, f"{type(exc).__name__}:{exc!s}"[:200]


def _pypdf_attempt_read_plain_content(reader: PyPDF2.PdfReader) -> Tuple[bool, str]:
    """Decrypt('') döndüsü başarısız görünürken ilk sayfanın düz metnine erişiliyorsa şifreyi gerektiren arayüz kapansın."""
    try:
        n = len(reader.pages)
        if n == 0:
            return False, "zero_pages_ambiguous"
        p0 = reader.pages[0]
        text = ""
        extract = getattr(p0, "extract_text", None)
        if callable(extract):
            text = extract() or ""
        stripped = (text or "").strip()
        snippet = stripped[:240] if stripped else "(empty_extract)"
        if not stripped:
            return False, "extract_blank_or_ws_only_under_encrypt_flag"
        return True, f"extract_text_ok_len_{len(snippet)}:{snippet[:48]!r}"


    except Exception as e:
        return False, f"{type(e).__name__}:{e!s}"[:220]


def _fallback_classify_via_pymupdf(
    pdf_path: str,
    diagnostics: Dict[str, Any],
    pdf_read_error: BaseException,
) -> Tuple[bool, str]:
    """
    PyPDF okuyamazsa bile PyMuPDF açabiliyorsa güvenilir sayfa kullanıcıya 'parolasız kullanılabilir' diye bildirilir.
    """
    needs, pym = pymupdf_requires_non_empty_password(pdf_path)
    diagnostics["engines"]["fallback"] = {
        "from_pypdf_error": True,
        "pypdf_error_type": type(pdf_read_error).__name__,
    }
    diagnostics["engines"]["pymupdf"] = pym
    if needs:
        logger.warning("[pdf-password] PyPDF unreadable and pymupdf still locked: %s", diagnostics)
        return True, "pypdf_failed_pymupdf_locked"
    logger.info("[pdf-password] PyPDF unreadable ama pymupdf parolasız: %s", diagnostics)
    return False, "pypdf_fatal_pymupdf_ok"


def is_pdf_encrypted(pdf_path: str) -> bool:
    """Kullanıcıdan (boş olmayan) açılış parolası istenmesi gerekiyorsa True."""
    req, _ = classify_pdf_password_requirement(pdf_path)
    return req


def validate_pdf_password(pdf_path: str, password: str) -> bool:
    """Parola PDF'i açabiliyorsa True döndürür; şifresiz dosyada her zaman True.
    Şifreli dosyada yalnızca verilen parola denenir (yanlış girişte boş parolaya düşülmez)."""
    basename = os.path.basename(pdf_path)
    try:
        with open(pdf_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh, strict=False)
            if not reader.is_encrypted:
                return True
            return bool(reader.decrypt((password or "").strip()))
    except Exception as e:
        raise Exception(_pdf_open_user_message("validate_pdf_password", e, basename)) from e


def get_num_pages(pdf_path: str, password: Optional[str] = None) -> int:
    """PDF içindeki sayfa sayısını döndürür; şifreliyse parola ile açar.
    Bölme ve önizleme akışları sayfa sınırı bilmek zorundadır.
    pypdf başarısız olursa PyMuPDF ile sayfa sayısı denenir."""
    try:
        reader = _open_pdf_reader(pdf_path, password=password, context="Sayfa bilgisi okuma")
        return len(reader.pages)
    except Exception as e:
        try:
            return _pymupdf_page_count(pdf_path, password=password)
        except Exception as e2:
            raise Exception(
                f"PDF sayfa sayısı okunamadı — pypdf: {e} | PyMuPDF: {e2}",
            ) from e2


def _word_to_pdf_win32com(docx_path: str, pdf_path: str) -> None:
    """
    Microsoft Word COM (pywin32) ile PDF dışa aktarır.
    Arka plan iş parçacığında çağrılıyorsa CoInitialize gerekir.
    """
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as e:
        raise ImportError("pywin32 paketi yüklü değil.") from e

    pythoncom.CoInitialize()
    word = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        try:
            word.DisplayAlerts = 0
        except Exception:
            pass
        doc_abs = os.path.abspath(docx_path)
        out_abs = os.path.abspath(pdf_path)
        doc = word.Documents.Open(doc_abs, ReadOnly=True)
        # Word sabiti: wdExportFormatPDF = 17 (sabit sayı COM API ile uyumludur).
        doc.ExportAsFixedFormat(OutputFileName=out_abs, ExportFormat=17, OpenAfterExport=False)
        doc.Close(SaveChanges=False)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def word_to_pdf(docx_path: str, pdf_path: str, progress_callback=None) -> bool:
    """
    Word belgesini PDF'e çevirir.

    Windows: Microsoft Word gerekir. Önce docx2pdf, yoksa pywin32 (COM) denenir.
    macOS: Microsoft Word + docx2pdf.
    """
    if sys.platform not in ("win32", "darwin"):
        raise Exception(
            "Word'den PDF dönüşümü şu an yalnızca Windows veya macOS üzerinde "
            "Microsoft Word yüklüyken desteklenir."
        )
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {docx_path}")
    ext = os.path.splitext(docx_path)[1].lower()
    if ext not in (".docx", ".doc"):
        raise ValueError("Yalnızca .docx veya .doc dosyaları desteklenir.")

    pdf_path = os.path.abspath(pdf_path)
    out_dir = os.path.dirname(pdf_path)
    if out_dir and not os.path.isdir(out_dir):
        raise FileNotFoundError(f"Hedef klasör bulunamadı: {out_dir}")

    if progress_callback:
        progress_callback(0, 2, "Microsoft Word ile PDF oluşturuluyor...")

    last_err: Optional[Exception] = None
    if sys.platform == "win32":
        try:
            from docx2pdf import convert

            try:
                convert(docx_path, pdf_path)
            except Exception as e:
                last_err = e
                try:
                    _word_to_pdf_win32com(docx_path, pdf_path)
                except Exception as e2:
                    raise Exception(
                        f"Word -> PDF Hatası (docx2pdf): {last_err}\n"
                        f"Yedek COM yolu da başarısız: {e2}"
                    ) from e2
        except ImportError:
            try:
                _word_to_pdf_win32com(docx_path, pdf_path)
            except ImportError:
                raise Exception(
                    "Word'den PDF için paket gerekli. PowerShell veya CMD'de (uygulamayı çalıştırdığın Python ile):\n\n"
                    "  python -m pip install docx2pdf\n\n"
                    "veya sadece COM yolu için:\n\n"
                    "  python -m pip install pywin32\n\n"
                    "Microsoft Word kurulu olmalıdır."
                ) from None
            except Exception as e:
                raise Exception(f"Word -> PDF Hatası: {e}") from e
    else:
        try:
            from docx2pdf import convert

            convert(docx_path, pdf_path)
        except ImportError as e:
            raise Exception(
                "Word'den PDF için 'docx2pdf' paketi gerekli. Kurulum:\n"
                "python -m pip install docx2pdf\n\n"
                "Microsoft Word for Mac kurulu olmalıdır."
            ) from e
        except Exception as e:
            raise Exception(f"Word -> PDF Hatası: {e}") from e

    if not os.path.isfile(pdf_path):
        raise Exception(
            "PDF dosyası oluşmadı. Word kapalı olsun, dosya başka programda açık olmasın "
            "veya Word güvenlik uyarısı engellenmiş olabilir."
        )

    if progress_callback:
        progress_callback(2, 2, "Tamamlandı")
    return True


def _fitz_pdf_text_stats(pdf_path: str, password: Optional[str]) -> Tuple[int, int]:
    """Sayfa sayısı ve çıkarılabilir metin uzunluğu (taranmış PDF ayırt etmek için)."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        if doc.needs_pass:
            if not password:
                raise Exception("Şifre gerekli")
            if not doc.authenticate(password):
                raise Exception("Girilen şifre hatalı")
        n = len(doc)
        total = 0
        for i in range(n):
            total += len(doc.load_page(i).get_text())
        return n, total
    finally:
        doc.close()


def _is_scanned_pdf(pdf_path: str, password: Optional[str] = None) -> bool:
    """
    Taranmış PDF tespiti: sayfaların büyük çoğunluğu görüntü blokları içeriyorsa
    ve metin çok azsa taranmış kabul et.

    Yöntem: fitz ile her sayfa blok analizi yap.
      - Görüntü bloğu var AND metin bloğu çok az → taranmış sayfa
      - Sayfaların %50+ taranmışsa → taranmış PDF
    """
    import fitz

    doc = fitz.open(pdf_path)
    try:
        if doc.needs_pass:
            doc.authenticate((password or "").strip())
        n = len(doc)
        if n == 0:
            return False
        scanned_pages = 0
        for i in range(n):
            page = doc[i]
            blocks = page.get_text("dict")["blocks"]
            img_blocks = sum(1 for b in blocks if b.get("type") == 1)
            # Metin karakteri sayısı (gerçek metin, boşluk sayılmaz)
            text_chars = sum(
                len(span["text"].strip())
                for b in blocks if b.get("type") == 0
                for line in b.get("lines", [])
                for span in line.get("spans", [])
            )
            # Sayfa taranmış mı: görüntü var AND metin çok az (sayfa başına <30 karakter)
            if img_blocks >= 1 and text_chars < 30:
                scanned_pages += 1
        return scanned_pages >= max(1, n * 0.5)
    finally:
        doc.close()


def _preprocess_ocr_image(img: "Image.Image") -> "Image.Image":
    """OCR doğruluğunu artırmak için görüntüyü ön işle: gri, kontrast, gürültü azaltma."""
    from PIL import ImageEnhance, ImageFilter

    # Gri tonlamaya çevir
    if img.mode != "L":
        img = img.convert("L")
    # Kontrast artır
    img = ImageEnhance.Contrast(img).enhance(1.8)
    # Keskinleştir
    img = img.filter(ImageFilter.SHARPEN)
    # Hafif gürültü azaltma
    img = img.filter(ImageFilter.MedianFilter(size=3))
    return img


def _ocr_page_to_blocks(img: "Image.Image", lang: str) -> List[Dict]:
    """
    Tesseract TSV (image_to_data) çıktısından kelime kutularını parse eder.
    Her blok: {'text', 'left', 'top', 'width', 'height', 'conf', 'block_num', 'par_num', 'line_num'}
    """
    from pytesseract import Output

    config = "--oem 3 --psm 6"
    data = pytesseract.image_to_data(img, lang=lang, config=config, output_type=Output.DICT)
    words = []
    n = len(data["text"])
    for i in range(n):
        txt = str(data["text"][i]).strip()
        conf = int(data["conf"][i])
        if not txt or conf < 20:
            continue
        words.append({
            "text":      txt,
            "left":      int(data["left"][i]),
            "top":       int(data["top"][i]),
            "width":     int(data["width"][i]),
            "height":    int(data["height"][i]),
            "conf":      conf,
            "block_num": int(data["block_num"][i]),
            "par_num":   int(data["par_num"][i]),
            "line_num":  int(data["line_num"][i]),
        })
    return words


def _words_to_paragraphs(words: List[Dict]) -> List[str]:
    """
    Kelime kutularını (block_num, par_num, line_num) üçlüsüne göre
    mantıklı paragraflara dönüştürür. Her blok ayrı paragraf.
    """
    if not words:
        return []

    # (block, par, line) → [words]
    lines: Dict[tuple, List[str]] = {}
    for w in words:
        key = (w["block_num"], w["par_num"], w["line_num"])
        lines.setdefault(key, []).append(w["text"])

    # Satırları (block, par) grubuna göre birleştir
    paras: Dict[tuple, List[str]] = {}
    for (b, p, _l), ws in sorted(lines.items()):
        key = (b, p)
        paras.setdefault(key, []).append(" ".join(ws))

    result = []
    for key in sorted(paras.keys()):
        para_text = " ".join(paras[key]).strip()
        para_text = _polish_tesseract_output(para_text)
        if para_text:
            result.append(para_text)
    return result


def _is_heading_candidate(text: str) -> bool:
    """Büyük harfli, kısa ve nokta/virgül içermeyen satırları başlık adayı say."""
    t = text.strip()
    if not t or len(t) > 120:
        return False
    words = t.split()
    if len(words) > 10:
        return False
    upper_ratio = sum(1 for c in t if c.isupper()) / max(len(t), 1)
    has_punct = any(c in t for c in ".,;:?!")
    return upper_ratio > 0.6 and not has_punct


def _build_docx_from_paragraphs(dw: "Document", paragraphs: List[str], page_num: int, total_pages: int) -> None:
    """Paragraph listesini Word belgesine yazar; başlık tespiti ve sayfa sonu ekler."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if page_num > 1:
        dw.add_page_break()

    if total_pages > 1:
        # Sayfa numarası başlığı — küçük, gri, ayırıcı işlevi
        h = dw.add_heading(f"— Sayfa {page_num} —", level=3)
        h.runs[0].font.size = Pt(9)

    if not paragraphs:
        p = dw.add_paragraph("(Bu sayfada metin okunamadı.)")
        p.runs[0].font.size = Pt(10)
        return

    for para_text in paragraphs:
        stripped = para_text.strip()
        if not stripped:
            continue
        if _is_heading_candidate(stripped):
            h = dw.add_heading(stripped, level=2)
            for run in h.runs:
                run.font.size = Pt(13)
        else:
            p = dw.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)


def _pdf_to_word_ocr_fitz(
    pdf_path: str,
    docx_path: str,
    password: Optional[str],
    progress_callback=None,
    *,
    ocr_matrix_scale: float = 2.5,
) -> None:
    """
    Taranmış / görsel PDF → düzenlenebilir Word.

    Strateji:
    1. Her sayfayı 300 DPI'ye yakın çözünürlükte render et
    2. Görüntü ön işleme: gri + kontrast + keskinleştirme
    3. Tesseract image_to_data (TSV) ile kelime bazlı bounding box al
    4. block_num/par_num/line_num gruplama ile paragraf yeniden oluştur
    5. Başlık tespiti + Word paragraph stilleri uygula
    """
    import fitz

    doc_pdf = fitz.open(pdf_path)
    try:
        if doc_pdf.needs_pass:
            if not password:
                raise Exception("Şifre gerekli")
            if not doc_pdf.authenticate(password):
                raise Exception("Girilen şifre hatalı")

        dw = Document()
        # Normal metin için varsayılan stil
        style = dw.styles["Normal"]
        style.font.name = "Calibri"
        from docx.shared import Pt
        style.font.size = Pt(11)

        n = len(doc_pdf)
        # Dil tespiti: ilk sayfada her iki dili de dene, hangisi daha fazla sonuç verirse onu kullan
        ocr_lang = "tur+eng"

        for i in range(n):
            if progress_callback:
                progress_callback(i + 1, max(n, 1), f"OCR işleniyor: sayfa {i + 1}/{n}")

            page = doc_pdf.load_page(i)
            mat = fitz.Matrix(ocr_matrix_scale, ocr_matrix_scale)
            pix = page.get_pixmap(matrix=mat, alpha=False, colorspace=fitz.csRGB)
            img = Image.open(io.BytesIO(pix.tobytes("png")))

            # Görüntü ön işleme
            img_proc = _preprocess_ocr_image(img)

            # Kelime bazlı OCR
            words = []
            try:
                words = _ocr_page_to_blocks(img_proc, ocr_lang)
            except Exception:
                try:
                    words = _ocr_page_to_blocks(img_proc, "eng")
                except Exception:
                    pass

            # Kelimeler yoksa ham metin dene
            if not words:
                try:
                    raw = pytesseract.image_to_string(img_proc, lang=ocr_lang, config="--oem 3 --psm 6")
                    paragraphs = [_polish_tesseract_output(b.strip()) for b in raw.split("\n\n") if b.strip()]
                except Exception:
                    paragraphs = []
            else:
                paragraphs = _words_to_paragraphs(words)

            _build_docx_from_paragraphs(dw, paragraphs, i + 1, n)

        dw.save(docx_path)
    finally:
        doc_pdf.close()


def pdf_to_word(
    pdf_path: str,
    docx_path: str,
    progress_callback=None,
    password: Optional[str] = None,
    *,
    reduced_quality: bool = False,
) -> bool:
    """
    PDF'i düzenlenebilir DOCX olarak dönüştürür.
    Metin katmanı çok zayıfsa Tesseract OCR ile metin üretir; aksi halde pdf2docx (yapısal) kullanılır.
    """
    try:
        if is_pdf_encrypted(pdf_path) and not password:
            raise Exception(f"PDF'ten Word'e dönüşüm için şifre gerekli: {os.path.basename(pdf_path)}")

        if progress_callback:
            progress_callback(0, 3, "PDF analiz ediliyor...")

        try:
            from pdf2docx import Converter
        except ImportError as e:
            raise Exception("pdf2docx yüklü değil.") from e

        # Taranmış PDF tespiti: fitz blok analizi ile görsel sayfa oranı kontrol edilir
        if _is_scanned_pdf(pdf_path, password):
            if progress_callback:
                progress_callback(1, 4, "Taranmış PDF tespit edildi — OCRmyPDF ile metin katmanı ekleniyor...")
            ocr_pdf_path = pdf_path + ".ocr_tmp.pdf"
            ocr_success = False
            try:
                import ocrmypdf
                ocrmypdf_kwargs = dict(
                    input_file=pdf_path,
                    output_file=ocr_pdf_path,
                    language="tur+eng",
                    deskew=True,
                    optimize=0,
                    progress_bar=False,
                    output_type="pdf",
                    skip_text=True,   # metin katmanı olan sayfaları atla, karışık PDF'lerde güvenli
                )
                if password:
                    ocrmypdf_kwargs["pdf_renderer"] = "hocr"
                ocrmypdf.ocr(**ocrmypdf_kwargs)
                ocr_success = os.path.isfile(ocr_pdf_path) and os.path.getsize(ocr_pdf_path) > 0
            except Exception as ocr_err:
                ocr_success = False
                print(f"[pdf_to_word] OCRmyPDF başarısız ({ocr_err}), Tesseract fallback'e geçiliyor")

            if ocr_success:
                if progress_callback:
                    progress_callback(2, 4, "Metin katmanı eklendi — Word'e dönüştürülüyor...")
                try:
                    _src = ocr_pdf_path
                    _pwd = None  # OCRmyPDF output is unencrypted
                    if os.path.isfile(docx_path):
                        try:
                            os.remove(docx_path)
                        except OSError:
                            pass
                    converter = Converter(_src, password=_pwd)
                    try:
                        converter.convert(
                            docx_path,
                            clip_image_res_ratio=2.0,
                            float_image_ignorable_gap=12.0,
                            line_overlap_threshold=0.9,
                            line_separate_threshold=6.0,
                        )
                    finally:
                        converter.close()
                    if progress_callback:
                        progress_callback(4, 4, "Word kaydediliyor...")
                    return True
                except Exception as conv_err:
                    print(f"[pdf_to_word] OCRmyPDF+pdf2docx başarısız ({conv_err}), Tesseract fallback")
                finally:
                    try:
                        if os.path.isfile(ocr_pdf_path):
                            os.remove(ocr_pdf_path)
                    except OSError:
                        pass

            # OCRmyPDF başarısız → eski Tesseract yolu
            if os.path.isfile(docx_path):
                try:
                    os.remove(docx_path)
                except OSError:
                    pass
            if progress_callback:
                progress_callback(2, 4, "Tesseract OCR ile metin çıkarılıyor...")
            ocr_scale = 1.38 if reduced_quality else 2.0
            _pdf_to_word_ocr_fitz(
                pdf_path,
                docx_path,
                password,
                progress_callback=progress_callback,
                ocr_matrix_scale=ocr_scale,
            )
            if progress_callback:
                progress_callback(4, 4, "Word kaydediliyor...")
            return True

        if progress_callback:
            progress_callback(0, 3, "Yapısal dönüşüm hazırlanıyor...")

        def _run_convert(use_ocr: int) -> None:
            converter = Converter(pdf_path, password=password)
            try:
                if progress_callback:
                    msg = (
                        "Taranmış sayfalar için OCR ile Word oluşturuluyor..."
                        if use_ocr
                        else "Sayfa düzeni korunarak Word oluşturuluyor..."
                    )
                    progress_callback(1, 3, msg)
                if os.path.isfile(docx_path):
                    try:
                        os.remove(docx_path)
                    except OSError:
                        pass
                converter.convert(
                    docx_path,
                    ocr=use_ocr,
                    clip_image_res_ratio=2.0,
                    float_image_ignorable_gap=12.0,
                    line_overlap_threshold=0.9,
                    line_separate_threshold=6.0,
                )
            finally:
                converter.close()

        try:
            _run_convert(0)
        except Exception as structural_error:
            try:
                _run_convert(1)
            except Exception:
                try:
                    if os.path.isfile(docx_path):
                        try:
                            os.remove(docx_path)
                        except OSError:
                            pass
                    if progress_callback:
                        progress_callback(2, 3, "Yapısal dönüşüm başarısız — OCR deneniyor...")
                    _fb_scale = 1.38 if reduced_quality else 2.0
                    _pdf_to_word_ocr_fitz(
                        pdf_path,
                        docx_path,
                        password,
                        progress_callback=progress_callback,
                        ocr_matrix_scale=_fb_scale,
                    )
                except Exception:
                    raise Exception(
                        "Bu PDF düzenlenebilir Word olarak dönüştürülemedi. "
                        "Tesseract kurulu olmalı; taranmış PDF’lerde OCR gerekir.\n\n"
                        f"Teknik ayrıntı: {structural_error}"
                    ) from structural_error

        if os.path.isfile(docx_path):
            if progress_callback:
                progress_callback(3, 3, "Word kaydediliyor...")
            return True

        raise Exception("Word dosyası oluşturulamadı.")

    except Exception as e:
        raise Exception(f"PDF -> Word Hatası: {e}")


def _safe_int(v, default=0) -> int:
    try:
        return int(v)
    except Exception:
        return default


def _safe_float(v, default=-1.0) -> float:
    try:
        return float(v)
    except Exception:
        return default


def _run_tesseract_image_to_data(page_proc: Image.Image) -> Dict[str, Any]:
    """
    Önce sadece Türkçe (Ü, Ğ vb. için daha iyi); yoksa tur+eng ile dener.
    """
    kwargs = dict(output_type=pytesseract.Output.DICT, config=_TESSERACT_CONFIG)
    try:
        return pytesseract.image_to_data(page_proc, lang="tur", **kwargs)
    except Exception:
        return pytesseract.image_to_data(page_proc, lang="tur+eng", **kwargs)


def _polish_tesseract_output(s: str) -> str:
    """
    Tesseract'ın sık yaptığı Türkçe / rakam hatalarını metin düzeyinde düzeltir.
    """
    if not s:
        return s
    # Ürün (Ü harfi bazen Uriin / Urin olarak gelir)
    s = re.sub(r"Uriin", "Ürün", s, flags=re.IGNORECASE)
    s = re.sub(r"\bUrin\b", "Ürün", s, flags=re.IGNORECASE)
    s = re.sub(r"\bUrun\b", "Ürün", s, flags=re.IGNORECASE)
    # Bitişik başlık
    s = re.sub(r"Ürün\s*Takip\s*Sistemi", "Ürün Takip Sistemi", s, flags=re.IGNORECASE)
    s = re.sub(r"ÜrünTakipSistemi", "Ürün Takip Sistemi", s, flags=re.IGNORECASE)
    # Rakamlar arası þ (thorn) -> 0 (ör. 735þ658411)
    s = re.sub(r"(\d)[þÞ](\d)", r"\g<1>0\g<2>", s)
    # Sadece rakam benzeri parçalarda kalan þ
    def fix_token(tok: str) -> str:
        if re.search(r"\d", tok) and re.search(r"[þÞ]", tok):
            return tok.replace("þ", "0").replace("Þ", "0")
        return tok

    parts = re.split(r"(\s+)", s)
    s = "".join(fix_token(p) if p.strip() else p for p in parts)
    return re.sub(r" {2,}", " ", s).strip()


def _finalize_line_text(s: str) -> str:
    return _polish_tesseract_output(_fix_glued_turkish_text(s))


def _fix_glued_turkish_text(s: str) -> str:
    """OCR/Tesseract bazen kelimeleri bitişik yazar (ör. FirmaAdi). Metin düzenlenebilir kalsın diye ayırır."""
    if not s:
        return s
    # küçük harf + büyük harf ayrımı (camelCase benzeri)
    s = re.sub(r"([a-zığüşöç])([A-ZĞÜŞİÖÇİ])", r"\1 \2", s)
    # rakam <-> harf
    s = re.sub(r"(\d)([A-Za-zğüşıöçĞÜŞİÖÇİ])", r"\1 \2", s)
    s = re.sub(r"([A-Za-zğüşıöçĞÜŞİÖÇİ])(\d)", r"\1 \2", s)
    # yaygın form etiketleri (ÜTS / bayilik ekranları)
    replacements = (
        ("FirmaAdi", "Firma Adı"),
        ("FirmaVergi", "Firma Vergi"),
        ("BayilikVeren", "Bayilik Veren"),
        ("BayilikAlan", "Bayilik Alan"),
        ("BayilikBaşvuru", "Bayilik Başvuru"),
        ("BayilikBasvuru", "Bayilik Başvuru"),
        ("BaşvuruTarihi", "Başvuru Tarihi"),
        ("BasvuruTarihi", "Başvuru Tarihi"),
        ("BaşlangıçTarihi", "Başlangıç Tarihi"),
        ("BaslangicTarihi", "Başlangıç Tarihi"),
        ("PlanlananBitiş", "Planlanan Bitiş"),
        ("PlanlananBitis", "Planlanan Bitiş"),
        ("BitişTarihi", "Bitiş Tarihi"),
        ("BitisTarihi", "Bitiş Tarihi"),
        ("KararTarihi", "Karar Tarihi"),
        ("İthalatBildirimi", "İthalat Bildirimi"),
        ("IthalatBildirimi", "İthalat Bildirimi"),
    )
    for a, b in replacements:
        s = s.replace(a, b)
    return re.sub(r" {2,}", " ", s).strip()


def _parse_ocr_words(ocr: Dict[str, Any]) -> List[Dict[str, Any]]:
    texts = ocr.get("text", [])
    if not texts:
        return []

    lefts = ocr.get("left", [])
    tops = ocr.get("top", [])
    widths = ocr.get("width", [])
    heights = ocr.get("height", [])
    block_nums = ocr.get("block_num", [0] * len(texts))
    par_nums = ocr.get("par_num", [0] * len(texts))
    line_nums = ocr.get("line_num", [0] * len(texts))
    confs = ocr.get("conf", [-1] * len(texts))

    words: List[Dict[str, Any]] = []
    n = len(texts)
    for i in range(n):
        t = (texts[i] or "").strip()
        if not t:
            continue
        conf = _safe_float(confs[i], default=-1.0)
        if conf < 0:
            conf = 0.0
        words.append(
            {
                "text": t,
                "left": _safe_int(lefts[i] if i < len(lefts) else 0),
                "top": _safe_int(tops[i] if i < len(tops) else 0),
                "width": _safe_int(widths[i] if i < len(widths) else 0),
                "height": _safe_int(heights[i] if i < len(heights) else 0),
                "block": _safe_int(block_nums[i] if i < len(block_nums) else 0),
                "par": _safe_int(par_nums[i] if i < len(par_nums) else 0),
                "line": _safe_int(line_nums[i] if i < len(line_nums) else 0),
                "conf": conf,
            }
        )
    return words


def _cluster_words_into_visual_lines(words: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    """Tüm blokları birleştirip y koordinatına göre görsel satırlara ayırır."""
    if not words:
        return []
    heights = sorted(w["height"] for w in words if w["height"] > 0)
    median_h = heights[len(heights) // 2] if heights else 14
    y_tol = max(10, int(median_h * 0.52))

    sorted_w = sorted(words, key=lambda w: (w["top"], w["left"]))
    lines: List[List[Dict[str, Any]]] = []
    current: List[Dict[str, Any]] = []
    anchor_y: float = 0.0

    for w in sorted_w:
        cy = float(w["top"] + max(1, w["height"]) / 2.0)
        if not current:
            current = [w]
            anchor_y = cy
            continue
        if abs(cy - anchor_y) <= y_tol:
            current.append(w)
            anchor_y = (anchor_y * (len(current) - 1) + cy) / len(current)
        else:
            current.sort(key=lambda x: x["left"])
            lines.append(current)
            current = [w]
            anchor_y = cy

    if current:
        current.sort(key=lambda x: x["left"])
        lines.append(current)
    return lines


def _join_words_in_segment(words: List[Dict[str, Any]]) -> str:
    if not words:
        return ""
    words = sorted(words, key=lambda x: x["left"])
    hlist = [w["height"] for w in words if w["height"] > 0]
    med_h = statistics.median(hlist) if hlist else 12.0
    space_threshold = max(2.0, med_h * 0.11)

    parts: List[str] = []
    for i, w in enumerate(words):
        if i == 0:
            parts.append(w["text"])
            continue
        prev = words[i - 1]
        gap = float(w["left"] - (prev["left"] + prev["width"]))
        sep = " " if gap > space_threshold else ""
        parts.append(sep + w["text"])
    raw = "".join(parts)
    return _finalize_line_text(raw.replace("  ", " ").strip())


def _split_line_into_word_segments(words: List[Dict[str, Any]], page_w: int) -> List[List[Dict[str, Any]]]:
    """Satır içinde etiket / değer gibi geniş boşluklardan kelime grupları üretir."""
    words = sorted(words, key=lambda x: x["left"])
    if len(words) <= 1:
        return [words]

    gaps: List[float] = []
    for i in range(1, len(words)):
        g = float(words[i]["left"] - (words[i - 1]["left"] + words[i - 1]["width"]))
        gaps.append(max(0.0, g))

    med_g = statistics.median(gaps) if gaps else 0.0
    # Form satırlarında etiket-değer arası genelde çok geniş; küçük iç boşluklara dokunma
    threshold = max(42.0, page_w * 0.052, med_g * 2.6 if med_g > 4.0 else 42.0)

    segments: List[List[Dict[str, Any]]] = []
    start = 0
    for i, g in enumerate(gaps):
        if g >= threshold:
            segments.append(words[start : i + 1])
            start = i + 1
    segments.append(words[start:])
    return [s for s in segments if s]


def _line_bbox(words: List[Dict[str, Any]]) -> Tuple[int, int]:
    left = min(w["left"] for w in words)
    right = max(w["left"] + w["width"] for w in words)
    return left, right


def _line_looks_centered(words: List[Dict[str, Any]], page_w: int) -> bool:
    if len(words) < 1 or page_w <= 0:
        return False
    left, right = _line_bbox(words)
    span = right - left
    mid = (left + right) / 2.0
    return abs(mid - page_w / 2.0) < page_w * 0.10 and span < page_w * 0.78


def _try_split_label_value_text(t: str) -> Optional[Tuple[str, str]]:
    """
    'Telefon: 0(212)...' veya 'Firma Adı: ...' gibi tek parçada kalan satırları [etiket, değer] ayırır.
    """
    t = (t or "").strip()
    if not t or t.lower().startswith("http"):
        return None
    first = t.split("\n", 1)[0].strip()
    if ":" not in first:
        return None
    idx = first.index(":")
    label = first[:idx].strip()
    rest = t[idx + 1 :].strip()
    if len(label) < 2 or len(label) > 58:
        return None
    # Çok geniş eşleşmeleri engelle (cümle içi iki nokta üst üste)
    if "\n" in label:
        return None
    return (label + ":", rest)


def _try_paragraph_to_two_column_row(text: str) -> Optional[Dict[str, Any]]:
    pair = _try_split_label_value_text(text)
    if not pair:
        return None
    label, value = pair
    return {"kind": "table_row", "cells": [label, value]}


def _looks_like_new_field_paragraph(text: str) -> bool:
    """Yeni bir form alanı satırı mı (sol tarafta etiket + iki nokta)?"""
    first = text.strip().split("\n", 1)[0].strip()
    if ":" not in first:
        return False
    idx = first.index(":")
    if idx > 58 or idx < 2:
        return False
    before = first[:idx].strip()
    if len(before) > 55:
        return False
    return True


def _looks_like_firma_adi_continuation(text: str) -> bool:
    """Firma adı değerinin ikinci satırı (SAN. TİC. LTD. ŞTİ. vb.)"""
    t = text.strip()
    if not t:
        return False
    if _looks_like_new_field_paragraph(t):
        return False
    if re.search(
        r"SAN\.\s*TİC\.|TİC\.\s*LTD|LTD\.?\s*ŞTİ\.?|ŞTİ\.?\s*\(|Üretici\s*/\s*İthalatçı",
        t,
        re.IGNORECASE,
    ):
        return True
    if t.startswith("(") and "Üretici" in t:
        return True
    return False


def _looks_like_phone_line(text: str) -> bool:
    """Telefon numarasına benzeyen kısa satır (OCR bazen etiketten ayırır)."""
    t = text.strip()
    if len(t) < 8:
        return False
    compact = re.sub(r"\s+", "", t)
    digits = sum(c.isdigit() for c in compact)
    if digits >= 10 and digits / max(1, len(compact)) > 0.5:
        return True
    if re.match(r"^0[\d\s\(\)\-]{8,}$", t.strip()):
        return True
    return False


def _should_merge_value_continuation(prev_row: Dict[str, Any], para_text: str) -> bool:
    """Önceki 2 sütunlu satırın sağ hücresine bu paragrafı eklemeli miyiz?"""
    if prev_row.get("kind") != "table_row":
        return False
    cells = prev_row.get("cells") or []
    if len(cells) != 2:
        return False
    left_l = cells[0].strip().lower()
    right = cells[1].strip()
    t = para_text.strip()
    if not t:
        return False
    if "bilgileri" in t.lower() and len(t) < 100 and "firma" not in t.lower():
        return False
    # Firma adı: uzun isim iki satıra bölünmüş
    if "firma adı" in left_l or "firma adi" in left_l:
        if _looks_like_firma_adi_continuation(t):
            return True
        if ":" not in t.split("\n")[0] and len(t) < 200:
            if right and not right.endswith(")"):
                if re.search(r"ŞTİ|TİC\.|LTD|SAN\.", t, re.IGNORECASE):
                    return True
    return False


def _normalize_single_cell_table_row(block: Dict[str, Any]) -> Dict[str, Any]:
    """Tek hücrede kalmış 'Etiket: değer' satırlarını ikiye böler."""
    if block.get("kind") != "table_row":
        return block
    cells = block.get("cells") or []
    if len(cells) != 1:
        return block
    pair = _try_split_label_value_text(cells[0])
    if not pair:
        return block
    return {"kind": "table_row", "cells": [pair[0], pair[1]]}


def _postprocess_form_layout_blocks(blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    - Paragraf olarak kalan 'Etiket: değer' satırlarını tablo satırına çevirir (Telefon hizası için).
    - Firma adı gibi alanlarda ikinci görsel satırı önceki satırın sağ hücresine birleştirir.
    """
    # A: paragrafları etiket/değer tablosuna çevir + tek hücreli tabloları böl
    pass1: List[Dict[str, Any]] = []
    for b in blocks:
        b = _normalize_single_cell_table_row(b)
        if b.get("kind") == "paragraph":
            conv = _try_paragraph_to_two_column_row(b.get("text", ""))
            if conv:
                pass1.append(conv)
                continue
        pass1.append(b)

    # B: devam satırlarını önceki satırın sağ hücresine yapıştır
    pass2: List[Dict[str, Any]] = []
    for b in pass1:
        if b.get("kind") == "paragraph" and pass2:
            prev = pass2[-1]
            if prev.get("kind") == "table_row" and len(prev.get("cells", [])) == 2:
                cells = prev["cells"]
                c0, c1 = cells[0].strip(), cells[1].strip()
                # Telefon etiketi tek satırda kaldıysa ve numara sonraki satırda geldiyse
                if (
                    c0.lower().startswith("telefon")
                    and ":" in c0
                    and _looks_like_phone_line(b.get("text", ""))
                    and len(c1) < 4
                ):
                    prev["cells"] = [cells[0], (c1 + " " + b["text"].strip()).strip()]
                    continue
            if _should_merge_value_continuation(prev, b.get("text", "")):
                cells = prev["cells"]
                prev["cells"] = [cells[0], (cells[1] + " " + b["text"].strip()).strip()]
                continue
        pass2.append(b)
    return pass2


def _is_section_heading_text(text: str) -> bool:
    t = text.strip()
    if len(t) > 90:
        return False
    if "Bilgileri" in t and len(t.split()) <= 8:
        return True
    if t.endswith(":") and len(t) < 50:
        return True
    return False


def _ocr_data_to_layout_blocks(ocr: Dict[str, Any], page_w: int) -> List[Dict[str, Any]]:
    """
    OCR çıktısını üstten alta sıralı blok listesine çevirir.
    - Etiket / değer: geniş boşlukla ayrılmış satırlar -> kenarlıksız tablo satırı
    - Tek sütun, ortada dar blok -> ortalanmış paragraf (başlık)
    """
    words = _parse_ocr_words(ocr)
    if not words:
        return []

    lines = _cluster_words_into_visual_lines(words)
    blocks: List[Dict[str, Any]] = []

    for line_words in lines:
        segs = _split_line_into_word_segments(line_words, page_w)
        texts = [_join_words_in_segment(s) for s in segs]
        texts = [t for t in texts if t.strip()]
        if not texts:
            continue

        if len(texts) == 1:
            text = texts[0]
            centered = _line_looks_centered(line_words, page_w)
            bold = _is_section_heading_text(text)
            blocks.append(
                {
                    "kind": "paragraph",
                    "text": text,
                    "centered": centered,
                    "bold": bold,
                }
            )
        else:
            # Çok sütun: üst bilgi (tarih | başlık | kullanıcı) veya yan yana alanlar
            blocks.append({"kind": "table_row", "cells": texts})

    return _postprocess_form_layout_blocks(blocks)


def _set_run_language_turkish(run) -> None:
    """Word yazım denetiminin Türkçe ile daha iyi çalışması için parça dilini ayarlar."""
    try:
        rpr = run._element.get_or_add_rPr()
        lang_el = OxmlElement("w:lang")
        lang_el.set(qn("w:val"), "tr-TR")
        lang_el.set(qn("w:eastAsia"), "tr-TR")
        rpr.append(lang_el)
    except Exception:
        pass


def _set_table_no_border(table) -> None:
    """Word tablosunda çizgileri kaldırır (düzen form görünümü)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tbl_pr.append(borders)


def _append_layout_block_to_document(doc: Document, block: Dict[str, Any]) -> None:
    if block.get("kind") == "paragraph":
        text = (block.get("text") or "").strip()
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(text)
        run.bold = bool(block.get("bold"))
        _set_run_language_turkish(run)
        if block.get("centered"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    if block.get("kind") == "table_row":
        cells = [c.strip() for c in block.get("cells") or [] if str(c).strip()]
        if not cells:
            return
        n = len(cells)
        tbl = doc.add_table(rows=1, cols=n)
        _set_table_no_border(tbl)
        row = tbl.rows[0]
        for j, ctext in enumerate(cells):
            cell = row.cells[j]
            cell.text = ""
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.left_indent = Pt(0)
            r = cp.add_run(ctext)
            r.bold = False
            _set_run_language_turkish(r)
        if n == 3:
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # İki sütun: sol etiket, sağ değer (tipik form)
        if n == 2:
            try:
                row.cells[0].width = Inches(2.25)
                row.cells[1].width = Inches(4.35)
            except Exception:
                pass
        return


def _fitz_open_for_tool(pdf_path: str, password: Optional[str] = None, context: str = "Bu işlem"):
    """PyMuPDF ile aç; boş kullanıcı parolası olan PDF'ler authenticate('') ile açılır."""
    import fitz

    basename = os.path.basename(pdf_path)
    doc = fitz.open(pdf_path)
    if doc.needs_pass:
        pwd = (password or "").strip()
        ok = False
        if pwd:
            ok = doc.authenticate(pwd)
        if not ok:
            ok = doc.authenticate("")
        if not ok:
            doc.close()
            if pwd:
                raise Exception(f"{context}: girilen PDF parolası hatalı ({basename}).")
            raise Exception(
                f"{context} için seçtiğiniz dosya şifreli: {basename}\n"
                "Lütfen dosya için şifre girin."
            )
    return doc


# --- 2. PDF BİRLEŞTİRME ---
def _pikepdf_open_for_merge(
    pdf_path: str,
    password: Optional[str],
    *,
    context: str = "PDF birleştirme",
):
    """pikepdf ile aç; PyMuPDF şifre mesajlarıyla uyumlu hatalar."""
    import pikepdf

    basename = os.path.basename(pdf_path)
    pwd_raw = (password or "").strip()

    def _raise_bad_password() -> None:
        raise Exception(f"{context}: girilen PDF parolası hatalı ({basename}).")

    def _raise_need_password() -> None:
        raise Exception(
            f"{context} için seçtiğiniz dosya şifreli: {basename}\n"
            "Lütfen dosya için şifre girin."
        )

    try:
        # pikepdf 10+: password must be str; None is rejected (TypeError).
        return pikepdf.open(pdf_path, password=pwd_raw)
    except pikepdf.PasswordError:
        if pwd_raw:
            _raise_bad_password()
        try:
            return pikepdf.open(pdf_path, password="")
        except pikepdf.PasswordError:
            _raise_need_password()


def _merge_pdfs_pikepdf(
    pdf_list: List[str],
    output_path: str,
    progress_callback,
    passwords: Optional[Dict[str, str]],
) -> None:
    """Çok sayfalı birleştirmede bellek dostu yol: pikepdf sayfa referansları ile birleştirir."""
    import gc
    import pikepdf

    passwords = passwords or {}
    # Birleştirme sırasında aralıklı ilerleme (sayfa başına callback RAM + kilidi şişirmez)
    page_chunk = int(os.environ.get("NB_MERGE_PAGE_CHUNK", "400"))
    page_chunk = max(50, min(page_chunk, 2000))

    total_pages = 0
    for pdf in pdf_list:
        if not os.path.isfile(pdf):
            raise FileNotFoundError(f"Birleştirilecek dosya bulunamadı: {pdf}")
        doc = _pikepdf_open_for_merge(
            pdf,
            passwords.get(pdf),
            context="PDF birleştirme",
        )
        try:
            total_pages += len(doc.pages)
        finally:
            doc.close()

    merged = pikepdf.Pdf.new()
    total = max(1, total_pages + 1)
    current_page = 0

    try:
        for pdf in pdf_list:
            src = _pikepdf_open_for_merge(
                pdf,
                passwords.get(pdf),
                context="PDF birleştirme",
            )
            try:
                base_name = os.path.basename(pdf)
                n = len(src.pages)
                pages = src.pages
                for start in range(0, n, page_chunk):
                    end = min(start + page_chunk, n)
                    merged.pages.extend(pages[start:end])
                    current_page += end - start
                    if progress_callback:
                        res = progress_callback(
                            current_page,
                            total,
                            f"{base_name} | Sayfa {end}/{n}",
                        )
                        if res is False:
                            raise Exception("İşlem iptal edildi.")
            finally:
                src.close()
                gc.collect()

        if progress_callback:
            progress_callback(total, total, "PDF yazılıyor...")
        merged.save(
            output_path,
            compress_streams=True,
            normalize_content=False,
            linearize=False,
        )
    finally:
        merged.close()


def _merge_pdfs_fitz_throttled(
    pdf_list: List[str],
    output_path: str,
    progress_callback,
    passwords: Optional[Dict[str, str]],
) -> None:
    """pikepdf uyumsuz PDF’ler için PyMuPDF geri dönüşü; ilerleme güncellemesi seyreltildi."""
    import fitz

    passwords = passwords or {}
    progress_every = int(os.environ.get("NB_MERGE_FITZ_PROGRESS_EVERY", "96"))
    progress_every = max(16, min(progress_every, 2000))

    total_pages = 0
    for pdf in pdf_list:
        if not os.path.isfile(pdf):
            raise FileNotFoundError(f"Birleştirilecek dosya bulunamadı: {pdf}")
        doc = _fitz_open_for_tool(pdf, passwords.get(pdf), context="PDF birleştirme")
        try:
            total_pages += doc.page_count
        finally:
            doc.close()

    merged = fitz.open()
    try:
        current_page = 0
        total = max(1, total_pages + 1)
        for pdf in pdf_list:
            src = _fitz_open_for_tool(pdf, passwords.get(pdf), context="PDF birleştirme")
            try:
                base_name = os.path.basename(pdf)
                n = src.page_count
                merged.insert_pdf(src)
                for page_idx in range(1, n + 1):
                    current_page += 1
                    at_last = page_idx == n
                    if progress_callback and (
                        at_last
                        or page_idx == 1
                        or (page_idx % progress_every == 0)
                    ):
                        res = progress_callback(
                            current_page,
                            total,
                            f"{base_name} | Sayfa {page_idx}/{n}",
                        )
                        if res is False:
                            raise Exception("İşlem iptal edildi.")
            finally:
                src.close()
        if progress_callback:
            progress_callback(total, total, "PDF yazılıyor...")
        merged.save(output_path, garbage=4, deflate=True, linear=False)
    finally:
        merged.close()


def merge_pdfs(pdf_list: List[str], output_path: str, progress_callback=None, passwords: Optional[Dict[str, str]] = None) -> bool:
    """
    Birden fazla PDF dosyasını tek çıktıda birleştirir.

    Varsayılan olarak pikepdf kullanılır (on binlerce sayfada PyMuPDF’e göre çok daha düşük tepe bellek).
    Başarısız olursa PyMuPDF ile yeniden denenir.
    """
    try:
        try:
            _merge_pdfs_pikepdf(pdf_list, output_path, progress_callback, passwords)
        except FileNotFoundError:
            raise
        except Exception as first:
            logger.info(
                "merge_pdfs_pikepdf_failed_trying_fitz",
                extra={"reason": str(first)},
            )
            _merge_pdfs_fitz_throttled(pdf_list, output_path, progress_callback, passwords)
        return True
    except Exception as e:
        err_text = str(e)
        if "PyCryptodome is required for AES algorithm" in err_text:
            raise Exception(
                "Bazı PDF dosyaları AES şifreleme kullanıyor. Birleştirme için eksik paket bulundu:\n\n"
                "python -m pip install pycryptodome"
            ) from e
        raise Exception(f"Birleştirme Hatası: {e}")


# --- 3. SAYFA AYIKLA (tek PDF olarak) ---
def extract_pages(
    pdf_path: str,
    pages: List[int],
    output_path: str,
    password: Optional[str] = None,
    output_password: Optional[str] = None,
) -> bool:
    """
    pdf_path içinden verilen sayfaları (1 tabanlı liste) alır ve tek PDF olarak output_path'e yazar.
    PyMuPDF ile tek kaynak açılır; büyük belgelerde bellek kullanımı daha kontrollüdür.
    """
    import fitz

    try:
        doc = _fitz_open_for_tool(pdf_path, password, context="Sayfa ayıklama")
        try:
            num_pages = doc.page_count
            normalized = []
            for p in pages:
                if not isinstance(p, int):
                    raise ValueError(f"Sayfa numarası tam sayı olmalıdır: {p}")
                if p < 1 or p > num_pages:
                    raise ValueError(f"Geçersiz sayfa numarası: {p} (Dosya {num_pages} sayfa)")
                normalized.append(p)

            out = fitz.open()
            try:
                for p in normalized:
                    out.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
                out.save(output_path, garbage=4, deflate=True, linear=False)
            finally:
                out.close()
        finally:
            doc.close()
        _apply_output_pdf_password(output_path, output_password)
        return True
    except Exception as e:
        raise Exception(f"Ayıklama Hatası: {e}")


# --- 4. SAYFA AYIKLA (her sayfa ayrı dosya olarak) ---
def extract_pages_separate(
    pdf_path: str,
    pages: List[int],
    output_folder: str,
    password: Optional[str] = None,
    output_password: Optional[str] = None,
) -> List[str]:
    """
    Verilen sayfaları (1 tabanlı) ayırır; her birini output_folder içinde ayrı PDF olarak kaydeder.
    Kaynak tek kez açılır; çok sayfalı dosyalarda daha verimlidir.
    """
    import fitz

    try:
        if not os.path.isdir(output_folder):
            raise FileNotFoundError(f"Hedef klasör bulunamadı: {output_folder}")

        doc = _fitz_open_for_tool(pdf_path, password, context="Sayfa ayıklama")
        try:
            num_pages = doc.page_count
            normalized = []
            for p in pages:
                if not isinstance(p, int):
                    raise ValueError(f"Sayfa numarası tam sayı olmalıdır: {p}")
                if p < 1 or p > num_pages:
                    raise ValueError(f"Geçersiz sayfa numarası: {p} (Dosya {num_pages} sayfa)")
                normalized.append(p)

            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_paths: List[str] = []

            for p in normalized:
                one = fitz.open()
                try:
                    one.insert_pdf(doc, from_page=p - 1, to_page=p - 1)
                    out_name = f"{base_name}_page_{p}.pdf"
                    out_path = os.path.join(output_folder, out_name)
                    one.save(out_path, garbage=4, deflate=True, linear=False)
                finally:
                    one.close()
                _apply_output_pdf_password(out_path, output_password)
                output_paths.append(out_path)

            return output_paths
        finally:
            doc.close()

    except Exception as e:
        raise Exception(f"Ayrı Ayırma Hatası: {e}")


# --- 5. PDF METİN / TABLO -> EXCEL ---
def _sanitize_sheet_title(title: str) -> str:
    invalid = '[]:*?/\\'
    cleaned = "".join("_" if ch in invalid else ch for ch in title).strip()
    return (cleaned or "Sayfa")[:31]


def _pdf_tables_to_excel(pdf_path: str, xlsx_path: str, progress_callback=None, password: Optional[str] = None) -> bool:
    try:
        import pdfplumber
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        raise Exception(
            "Tablo koruma modu için 'pdfplumber' ve 'openpyxl' gerekli: python -m pip install pdfplumber openpyxl"
        ) from e

    try:
        wb = Workbook()
        default_ws = wb.active
        wb.remove(default_ws)

        if is_pdf_encrypted(pdf_path) and not password:
            raise Exception(f"PDF -> Excel için şifre gerekli: {os.path.basename(pdf_path)}")
        with pdfplumber.open(pdf_path, password=password) as pdf:
            total_pages = len(pdf.pages)
            border = Border(
                left=Side(style="thin", color="D7DEE8"),
                right=Side(style="thin", color="D7DEE8"),
                top=Side(style="thin", color="D7DEE8"),
                bottom=Side(style="thin", color="D7DEE8"),
            )
            title_fill = PatternFill("solid", fgColor="1F4E78")
            header_fill = PatternFill("solid", fgColor="DCE6F1")
            for i, page in enumerate(pdf.pages, start=1):
                if progress_callback:
                    progress_callback(i, max(1, total_pages), f"Tablo aranıyor: Sayfa {i}/{total_pages}")

                ws = wb.create_sheet(_sanitize_sheet_title(f"Sayfa {i}"))
                ws.sheet_view.showGridLines = False
                ws.freeze_panes = "A2"

                tables = page.extract_tables(
                    table_settings={
                        "vertical_strategy": "lines",
                        "horizontal_strategy": "lines",
                        "snap_tolerance": 4,
                        "join_tolerance": 4,
                        "intersection_tolerance": 6,
                    }
                ) or []
                if not tables:
                    tables = page.extract_tables(
                        table_settings={
                            "vertical_strategy": "text",
                            "horizontal_strategy": "text",
                            "text_x_tolerance": 2,
                            "text_y_tolerance": 2,
                        }
                    ) or []
                current_row = 1

                cleaned_tables = []
                for table in tables:
                    cleaned_rows = []
                    for row in table or []:
                        cells = ["" if cell is None else str(cell).strip() for cell in row]
                        if any(cell for cell in cells):
                            cleaned_rows.append(cells)
                    if cleaned_rows:
                        cleaned_tables.append(cleaned_rows)

                if cleaned_tables:
                    for table_index, table_rows in enumerate(cleaned_tables, start=1):
                        title_cell = ws.cell(row=current_row, column=1, value=f"Tablo {table_index}")
                        title_cell.font = Font(bold=True, color="FFFFFF")
                        title_cell.fill = title_fill
                        title_cell.alignment = Alignment(horizontal="left", vertical="center")
                        current_row += 1
                        max_cols = max(len(r) for r in table_rows)
                        col_max = [0] * max_cols
                        for row_offset, row in enumerate(table_rows, start=0):
                            padded = row + [""] * (max_cols - len(row))
                            for col_index, value in enumerate(padded, start=1):
                                cell = ws.cell(row=current_row, column=col_index, value=value)
                                cell.border = border
                                cell.alignment = Alignment(vertical="center", horizontal="left", wrap_text=True)
                                if row_offset == 0:
                                    cell.font = Font(bold=True)
                                    cell.fill = header_fill
                                col_max[col_index - 1] = max(col_max[col_index - 1], len(str(value or "")))
                            ws.row_dimensions[current_row].height = 22
                            current_row += 1
                        for col_index, length in enumerate(col_max, start=1):
                            ws.column_dimensions[get_column_letter(col_index)].width = min(40, max(12, length + 2))
                        current_row += 2
                else:
                    text = (page.extract_text() or "").strip()
                    ws.cell(row=1, column=1, value="Bu sayfada tablo bulunamadı.")
                    if text:
                        ws.cell(row=3, column=1, value="Algılanan metin")
                        for row_index, line in enumerate([ln.strip() for ln in text.splitlines() if ln.strip()], start=4):
                            ws.cell(row=row_index, column=1, value=line)

        if not wb.sheetnames:
            ws = wb.create_sheet("Sayfa 1")
            ws.cell(row=1, column=1, value="İçerik bulunamadı.")

        wb.save(xlsx_path)
        return True
    except Exception as e:
        raise Exception(f"PDF tablo -> Excel Hatası: {e}") from e


def pdf_text_to_excel(
    pdf_path: str,
    xlsx_path: str,
    progress_callback=None,
    preserve_tables: bool = False,
    password: Optional[str] = None,
) -> bool:
    """
    PDF'i Excel'e aktarır.
    preserve_tables=True ise önce tablo yapısını korumaya çalışır.
    Aksi halde sayfa/satır bazlı metin aktarımı yapar.
    """
    if preserve_tables:
        try:
            return _pdf_tables_to_excel(pdf_path, xlsx_path, progress_callback=progress_callback, password=password)
        except Exception as e:
            raise Exception(f"PDF -> Excel (tablo koruma) Hatası: {e}") from e

    try:
        from openpyxl import Workbook
    except ImportError as e:
        raise Exception("PDF -> Excel için 'openpyxl' gerekli: python -m pip install openpyxl") from e

    try:
        reader = _open_pdf_reader(pdf_path, password=password, context="PDF -> Excel")
        num = len(reader.pages)
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Metni"
        ws.append(["Sayfa", "Satır No", "Metin"])

        for i in range(num):
            if progress_callback:
                progress_callback(i + 1, max(1, num), f"Sayfa {i + 1}/{num}")
            page = reader.pages[i]
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            if not lines:
                ws.append([i + 1, 1, "(Bu sayfada metin çıkarılamadı)"])
            else:
                for j, line in enumerate(lines, start=1):
                    ws.append([i + 1, j, line])

        wb.save(xlsx_path)
        return True
    except Exception as e:
        raise Exception(f"PDF -> Excel Hatası: {e}")


def _excel_to_pdf_win32com(xlsx_path: str, pdf_path: str) -> None:
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as e:
        raise ImportError("pywin32 gerekli.") from e

    pythoncom.CoInitialize()
    xl = None
    try:
        xl = DispatchEx("Excel.Application")
        xl.Visible = False
        try:
            xl.DisplayAlerts = False
        except Exception:
            pass
        wb = xl.Workbooks.Open(os.path.abspath(xlsx_path), ReadOnly=True)
        # 0 = xlTypePDF
        wb.ExportAsFixedFormat(0, os.path.abspath(pdf_path))
        wb.Close(SaveChanges=False)
    finally:
        if xl is not None:
            try:
                xl.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _excel_to_pdf_reportlab(xlsx_path: str, pdf_path: str) -> None:
    try:
        from openpyxl import load_workbook
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
    except ImportError as e:
        raise ImportError("openpyxl ve reportlab gerekli.") from e

    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    ws = wb.active
    rows: List[List[str]] = []
    max_cols = 0
    for row in ws.iter_rows(values_only=True):
        cells = ["" if c is None else str(c) for c in row]
        while cells and cells[-1] == "":
            cells.pop()
        if cells:
            max_cols = max(max_cols, len(cells))
            rows.append(cells)
    wb.close()

    if not rows:
        rows = [["(Boş sayfa)"]]
        max_cols = 1

    for r in rows:
        while len(r) < max_cols:
            r.append("")

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    usable_w = landscape(A4)[0] - 24 * mm
    col_w = usable_w / max(1, max_cols)
    t = Table(rows, colWidths=[col_w] * max_cols, repeatRows=1 if len(rows) > 1 else 0)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3a86ff")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
            ]
        )
    )
    doc.build([t])


# --- 6. EXCEL -> PDF ---
def excel_to_pdf(xlsx_path: str, pdf_path: str, progress_callback=None) -> bool:
    """
    Excel'i PDF yapar.
    Windows: Excel COM (Office). Diğer / yedek: openpyxl + reportlab (ilk sayfa, tablo görünümü).
    """
    if not os.path.isfile(xlsx_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {xlsx_path}")
    ext = os.path.splitext(xlsx_path)[1].lower()
    if ext not in (".xlsx", ".xlsm", ".xltx", ".xltm"):
        raise ValueError("Desteklenen biçimler: .xlsx, .xlsm (eski .xls desteklenmez)")

    pdf_path = os.path.abspath(pdf_path)
    out_dir = os.path.dirname(pdf_path)
    if out_dir and not os.path.isdir(out_dir):
        raise FileNotFoundError(f"Hedef klasör bulunamadı: {out_dir}")

    if progress_callback:
        progress_callback(0, 2, "PDF oluşturuluyor...")

    if sys.platform == "win32":
        try:
            _excel_to_pdf_win32com(xlsx_path, pdf_path)
        except Exception:
            try:
                _excel_to_pdf_reportlab(xlsx_path, pdf_path)
            except Exception as e2:
                raise Exception(
                    "Excel -> PDF başarısız. Excel yüklüyse pywin32 kurun; değilse:\n"
                    "python -m pip install openpyxl reportlab\n\n"
                    f"Ayrıntı: {e2}"
                ) from e2
    else:
        try:
            _excel_to_pdf_reportlab(xlsx_path, pdf_path)
        except Exception as e:
            raise Exception(f"Excel -> PDF Hatası: {e}") from e

    if not os.path.isfile(pdf_path):
        raise Exception("PDF dosyası oluşmadı.")

    if progress_callback:
        progress_callback(2, 2, "Tamamlandı")
    return True


def _tool_subprocess_timeout_sec() -> int:
    return max(30, int(os.environ.get("NB_PDF_TOOL_TIMEOUT_SEC", "60")))


def _resolve_ghostscript_executable() -> Optional[str]:
    """Windows'ta önce ``gswin64c``; ortam: ``NB_GHOSTSCRIPT_EXE`` tam yol."""
    override = (os.environ.get("NB_GHOSTSCRIPT_EXE") or "").strip()
    if override and os.path.isfile(override):
        return override
    if platform.system() == "Windows":
        for cand in (shutil.which("gswin64c"), shutil.which("gswin32c"), shutil.which("gs")):
            if cand:
                return cand
        for pf_key in ("ProgramFiles", "ProgramFiles(x86)"):
            pf = os.environ.get(pf_key, "")
            if not pf:
                continue
            gs_root = os.path.join(pf, "gs")
            if not os.path.isdir(gs_root):
                continue
            for ver in sorted(os.listdir(gs_root), reverse=True):
                for exe_name in ("gswin64c.exe", "gswin32c.exe"):
                    p = os.path.join(gs_root, ver, "bin", exe_name)
                    if os.path.isfile(p):
                        return p
    w = shutil.which("gs")
    return w


def _ghostscript_compress_to_path(
    input_path: str,
    output_path: str,
    *,
    pdfsettings: str,
    password: str,
    timeout_sec: int,
) -> bool:
    """Ghostscript pdfwrite ile sıkıştırma; başarıda ``output_path`` yazar."""
    exe = _resolve_ghostscript_executable()
    if not exe:
        return False
    src_abs = os.path.abspath(input_path)
    out_abs = os.path.abspath(output_path)
    out_dir = os.path.dirname(out_abs) or "."
    fd, tmp_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    tmp_keep: Optional[str] = tmp_out
    try:
        dpi_map = {"/printer": 300, "/ebook": 150, "/screen": 96, "/prepress": 300}
        dpi = dpi_map.get(pdfsettings, 150)
        cmd: List[str] = [
            exe,
            "-sDEVICE=pdfwrite",
            "-dNOPAUSE",
            "-dBATCH",
            "-dQUIET",
            "-dSAFER",
            "-dCompatibilityLevel=1.5",
            f"-dPDFSETTINGS={pdfsettings}",
            "-dDetectDuplicateImages=true",
            "-dCompressFonts=true",
            "-dSubsetFonts=true",
            "-dDownsampleColorImages=true",
            "-dDownsampleGrayImages=true",
            "-dColorImageDownsampleType=/Bicubic",
            "-dGrayImageDownsampleType=/Bicubic",
            f"-dColorImageResolution={dpi}",
            f"-dGrayImageResolution={dpi}",
            "-dMonoImageResolution=300",
            "-dOptimize=true",
            "-dFastWebView=true",
            f"-sOutputFile={tmp_out}",
            src_abs,
        ]
        if (password or "").strip():
            cmd.insert(1, f"-sPDFPassword={password.strip()}")
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout_sec)
        if r.returncode != 0 or not os.path.isfile(tmp_out) or os.path.getsize(tmp_out) < 64:
            return False
        os.replace(tmp_out, out_abs)
        tmp_keep = None
        return True
    except (subprocess.TimeoutExpired, OSError):
        return False
    finally:
        if tmp_keep and os.path.isfile(tmp_keep):
            try:
                os.remove(tmp_keep)
            except OSError:
                pass


def _pdf_page_count(path: str, password: str = "") -> int:
    """PyMuPDF ile sayfa sayısını döndürür; hata olursa -1."""
    try:
        import fitz
        doc = fitz.open(path)
        if password and doc.needs_pass:
            doc.authenticate(password)
        n = doc.page_count
        doc.close()
        return n
    except Exception:
        return -1


def _fast_image_compress_pipeline(
    input_path: str, output_path: str, open_password: str, quality: str = "auto"
) -> bool:
    """pikepdf + Pillow ile hızlı görüntü çözünürlük düşürme + yeniden sıkıştırma.

    Tüm PDF dolaylı nesnelerini tarar (/Subtype /Image olanları sıkıştırır).
    page.images yerine doğrudan nesne tablosu kullanılır — form XObject içindeki
    görüntüler de dahil olmak üzere hiçbir görüntü atlanmaz.
    200 MB dosyada ~30-90 saniye. Döner: True = küçültme başarılı.

    Hedef uzun kenar (piksel):
      low    → 1200 px  (~100 DPI A4) + JPEG 50
      auto   → 1800 px  (~150 DPI A4) + JPEG 62
      medium → 1800 px  (~150 DPI A4) + JPEG 72
      high   → 2400 px  (~200 DPI A4) + JPEG 82
    """
    import io
    import pikepdf
    from PIL import Image

    cfg = {
        "low":    {"max_px": 900,  "jpeg_q": 40},
        "auto":   {"max_px": 1400, "jpeg_q": 52},
        "medium": {"max_px": 1600, "jpeg_q": 65},
        "high":   {"max_px": 2200, "jpeg_q": 78},
    }
    c = cfg.get(quality, cfg["auto"])
    max_px: int = c["max_px"]
    jpeg_q: int = c["jpeg_q"]

    out_dir = os.path.dirname(output_path) or None
    fd, tmp_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    tmp_keep: Optional[str] = tmp_out

    try:
        open_kw: dict = {}
        if open_password:
            open_kw["password"] = open_password

        pdf = pikepdf.open(input_path, allow_overwriting_input=False, **open_kw)
        replaced = 0

        # Tüm dolaylı nesneleri tara — page.images sadece sayfa kaynakları sözlüğündeki
        # adlandırılmış görüntüleri döndürür; form XObject ve paylaşılan nesneler atlanabilir.
        for obj in pdf.objects:
            try:
                if not isinstance(obj, pikepdf.Stream):
                    continue
                if obj.get("/Subtype") != pikepdf.Name("/Image"):
                    continue

                # Maske görüntülerini atla (genellikle 1-bit, JPEG'e uygun değil)
                color_space = obj.get("/ColorSpace")
                if color_space == pikepdf.Name("/DeviceGray"):
                    bits = obj.get("/BitsPerComponent")
                    if bits is not None and int(bits) == 1:
                        continue

                w = int(obj.get("/Width", 0))
                h = int(obj.get("/Height", 0))
                if w < 8 or h < 8:
                    continue

                # Ham stream boyutunu kontrol et — 5 KB altı küçük logolar/ikonları atla
                try:
                    raw = obj.read_raw_bytes()
                except Exception:
                    continue
                if len(raw) < 5 * 1024:
                    continue

                # Pillow'a çevir
                try:
                    pdfimg = pikepdf.PdfImage(obj)
                    pil_img: Image.Image = pdfimg.as_pil_image()
                except Exception:
                    continue

                # Mod dönüşümü: JPEG alpha desteklemez
                if pil_img.mode == "P":
                    pil_img = pil_img.convert("RGBA")
                if pil_img.mode in ("RGBA", "LA"):
                    bg = Image.new("RGB", pil_img.size, (255, 255, 255))
                    bg.paste(pil_img, mask=pil_img.split()[-1])
                    pil_img = bg
                elif pil_img.mode not in ("RGB", "L"):
                    pil_img = pil_img.convert("RGB")

                # Çözünürlük düşürme: en uzun kenar max_px'i geçiyorsa küçült
                long_side = max(w, h)
                if long_side > max_px:
                    scale = max_px / long_side
                    new_w = max(1, int(w * scale))
                    new_h = max(1, int(h * scale))
                    pil_img = pil_img.resize((new_w, new_h), Image.LANCZOS)

                # JPEG olarak kodla
                buf = io.BytesIO()
                pil_img.save(buf, format="JPEG", quality=jpeg_q, optimize=True, progressive=True)
                new_bytes = buf.getvalue()

                # Orijinalden büyükse atla
                if len(new_bytes) >= len(raw):
                    continue

                # Stream'i yeni JPEG verisiyle güncelle
                obj.write(new_bytes, filter=pikepdf.Name("/DCTDecode"))
                # ColorSpace ve BitsPerComponent'i JPEG ile uyumlu hale getir
                if pil_img.mode == "L":
                    obj["/ColorSpace"] = pikepdf.Name("/DeviceGray")
                else:
                    obj["/ColorSpace"] = pikepdf.Name("/DeviceRGB")
                obj["/BitsPerComponent"] = pikepdf.objects.Integer(8)
                obj["/Width"] = pikepdf.objects.Integer(pil_img.width)
                obj["/Height"] = pikepdf.objects.Integer(pil_img.height)
                # Önceki filter array'ini temizle
                for k in ("/DecodeParms", "/Decode"):
                    if k in obj:
                        del obj[k]
                replaced += 1

            except Exception:
                continue

        # Thumbnail'ları temizle
        try:
            for page in pdf.pages:
                if "/Thumb" in page:
                    del page["/Thumb"]
        except Exception:
            pass

        pdf.save(
            tmp_out,
            compress_streams=True,
            recompress_flate=True,
            object_stream_mode=pikepdf.ObjectStreamMode.generate,
            linearize=True,
        )
        pdf.close()

        in_size = os.path.getsize(input_path)
        out_size = os.path.getsize(tmp_out)
        if out_size < in_size:
            os.replace(tmp_out, output_path)
            tmp_keep = None
            return True
        # Sıkıştırma kazanımı yoksa bile stream-optimize edilmiş sürümü kullan
        # (en azından orijinal kadar iyi)
        shutil.copy2(input_path, output_path)
        return False

    except Exception as e:
        logger.warning("_fast_image_compress_pipeline failed: %s", e)
        shutil.copy2(input_path, output_path)
        return False
    finally:
        if tmp_keep and os.path.isfile(tmp_keep):
            try:
                os.remove(tmp_keep)
            except OSError:
                pass


def _fitz_render_compress(
    input_path: str, output_path: str, open_password: str, quality: str = "auto"
) -> bool:
    """PyMuPDF ile sayfa bazlı render + JPEG yeniden kodlama (Ghostscript alternatifi).

    Her sayfayı hedef DPI'de piksel olarak render eder, ardından yeni bir PDF içine
    JPEG stream olarak yerleştirir. Metin aranabilirliği kaybolur ama görüntü ağırlıklı
    (tarama) PDF'lerde en agresif boyut küçültmeyi sağlar.
    """
    import fitz
    from PIL import Image as PilImage

    dpi_map   = {"low": 96,  "auto": 120, "medium": 150, "high": 200}
    jpegq_map = {"low": 35,  "auto": 50,  "medium": 62,  "high": 75}
    dpi   = dpi_map.get(quality, 120)
    jpegq = jpegq_map.get(quality, 50)

    out_dir = os.path.dirname(output_path) or None
    fd, tmp_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    tmp_keep: Optional[str] = tmp_out

    try:
        src = fitz.open(input_path)
        if open_password and src.needs_pass:
            src.authenticate(open_password)

        out_doc = fitz.open()
        scale = dpi / 72.0
        mat = fitz.Matrix(scale, scale)

        for page in src:
            pix = page.get_pixmap(matrix=mat, alpha=False, colorspace=fitz.csRGB)
            # Pillow ile JPEG olarak yeniden kodla (kalite kontrolü için)
            pil_img = PilImage.frombytes("RGB", (pix.width, pix.height), pix.samples)
            buf = io.BytesIO()
            pil_img.save(buf, format="JPEG", quality=jpegq, optimize=True, progressive=True)
            jpeg_bytes = buf.getvalue()

            # Orijinal sayfa boyutlarını koru (pt cinsinden)
            rect = page.rect
            new_page = out_doc.new_page(width=rect.width, height=rect.height)
            new_page.insert_image(rect, stream=jpeg_bytes)

        src.close()
        out_doc.save(
            tmp_out,
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
        )
        out_doc.close()

        in_size  = os.path.getsize(input_path)
        out_size = os.path.getsize(tmp_out)
        if out_size < in_size:
            os.replace(tmp_out, output_path)
            tmp_keep = None
            return True
        shutil.copy2(input_path, output_path)
        return False

    except Exception as e:
        logger.warning("_fitz_render_compress failed: %s", e)
        shutil.copy2(input_path, output_path)
        return False
    finally:
        if tmp_keep and os.path.isfile(tmp_keep):
            try:
                os.remove(tmp_keep)
            except OSError:
                pass


def _legacy_compress_pipeline(input_path: str, output_path: str, open_password: str) -> None:
    """PyMuPDF stream-sıkıştırma yedek boru hattı (görüntü yeniden örnekleme yapmaz)."""
    import fitz

    out_dir = os.path.dirname(output_path) or None
    fd, tmp_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        doc = fitz.open(input_path)
        if open_password and doc.needs_pass:
            doc.authenticate(open_password)
        doc.save(
            tmp_out,
            garbage=4,
            deflate=True,
            deflate_images=True,
            deflate_fonts=True,
            clean=True,
            linear=False,
        )
        doc.close()
        if os.path.getsize(tmp_out) < os.path.getsize(input_path):
            os.replace(tmp_out, output_path)
            tmp_out = None
        else:
            shutil.copy2(input_path, output_path)
    except Exception:
        raise
    finally:
        if tmp_out and os.path.isfile(tmp_out):
            try:
                os.remove(tmp_out)
            except OSError:
                pass


def _fmt_mb(n: int) -> str:
    return f"{n / (1024 * 1024):.1f} MB"


# --- 7. PDF SIKIŞTIRMA ---
# Kaliteye göre GS preset ve kazanım eşiği
_COMPRESS_GS_SETTINGS = {
    "low":    "/screen",
    "auto":   "/ebook",
    "medium": "/ebook",
    "high":   "/printer",
}
# pikepdf sonrası bu oran altında kazanım varsa GS devreye girer (0.45 = %45)
_COMPRESS_GS_FALLBACK_THRESHOLD = 0.40


def compress_pdf(input_path: str, output_path: str, progress_callback=None, password: Optional[str] = None, quality: str = "auto") -> bool:
    """İki aşamalı sıkıştırma: pikepdf+Pillow görüntü yeniden örnekleme + gerekirse Ghostscript.

    Strateji:
      1. pikepdf + Pillow: görüntüleri hedef çözünürlüğe düşür, JPEG ile yeniden kodla.
      2. pikepdf kazanımı hedef eşiğin altındaysa Ghostscript preset ile ikinci geçiş.
         low    → /screen   (en agresif)
         auto   → /ebook    (150 DPI, dengeli hız/kalite)
         medium → /ebook
         high   → /printer  (300 DPI, en kaliteli)
    Her iki aşamada da en küçük sonuç seçilir.
    """
    try:
        if progress_callback:
            progress_callback(0, 2, "PDF sıkıştırılıyor...")
        if is_pdf_encrypted(input_path) and not password:
            raise Exception(f"PDF sıkıştırma için şifre gerekli: {os.path.basename(input_path)}")
        open_password = (password or "").strip()
        in_size = os.path.getsize(input_path)
        timeout_sec = _tool_subprocess_timeout_sec()

        out_dir = os.path.dirname(output_path) or None
        pike_keep: Optional[str] = None
        gs_keep: Optional[str] = None

        # --- 1. pikepdf + Pillow görüntü yeniden sıkıştırma ---
        fd1, pike_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
        os.close(fd1)
        pike_keep = pike_out
        try:
            _fast_image_compress_pipeline(input_path, pike_out, open_password, quality)
            pike_size = os.path.getsize(pike_out)
            pike_ratio = 1.0 - pike_size / in_size  # örn. 0.42 = %42 küçülme
        except Exception:
            pike_size = in_size
            pike_ratio = 0.0

        if progress_callback:
            progress_callback(1, 2, "Görüntüler işlendi, akışlar optimize ediliyor...")

        # --- 2. GS ikinci geçiş: pikepdf yeterince sıkıştıramadıysa veya "high" ise ---
        best_path: str = pike_out
        best_size: int = pike_size
        need_second = False

        # --- 2. İkinci geçiş: önce GS dene, yoksa fitz render ---
        # pikepdf yeterli kazanım sağlayamadıysa veya "high" kaliteyse ikinci geçiş zorunlu
        need_second = (quality == "high") or (pike_ratio < _COMPRESS_GS_FALLBACK_THRESHOLD)

        if need_second:
            gs_preset = _COMPRESS_GS_SETTINGS.get(quality, "/ebook")
            gs_exe = _resolve_ghostscript_executable()

            fd2, second_out = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
            os.close(fd2)
            gs_keep = second_out
            # İkinci geçişin girişi: pikepdf küçülttüyse onu kullan
            second_input = pike_out if pike_size < in_size else input_path
            second_ok = False

            if gs_exe:
                # GS kuruluysa tercih et
                second_ok = _ghostscript_compress_to_path(
                    second_input, second_out,
                    pdfsettings=gs_preset,
                    password=open_password,
                    timeout_sec=timeout_sec,
                )
            else:
                # GS yok → fitz sayfa render pipeline (metin aranabilirliği kaybolur
                # ama görüntü ağırlıklı/tarama PDF'lerde en etkili sıkıştırma)
                second_ok = _fitz_render_compress(second_input, second_out, open_password, quality)

            if second_ok:
                second_size = os.path.getsize(second_out)
                try:
                    exp_pages = _pdf_page_count(second_input)
                    got_pages = _pdf_page_count(second_out)
                    pages_ok = (exp_pages <= 0 or got_pages == exp_pages)
                except Exception:
                    pages_ok = True
                if pages_ok and second_size < best_size:
                    best_path = second_out
                    best_size = second_size

        # --- En küçük sonucu çıktıya yaz ---
        if best_size < in_size:
            shutil.copy2(best_path, output_path)
        else:
            shutil.copy2(input_path, output_path)

        saved_pct = max(0, round((1 - best_size / in_size) * 100, 1))
        logger.info(
            "compress_pdf done: %s → %s (%.1f%% saved, quality=%s, gs=%s)",
            _fmt_mb(in_size), _fmt_mb(best_size), saved_pct, quality, need_second,
        )

        if progress_callback:
            progress_callback(2, 2, "Tamamlandı")
        return True
    except Exception as e:
        raise Exception(f"PDF Sıkıştırma Hatası: {e}")
    finally:
        for _tmp in (pike_keep, gs_keep):
            if _tmp and os.path.isfile(_tmp):
                try:
                    os.remove(_tmp)
                except OSError:
                    pass


# --- 8. PDF ŞİFRELEME ---
def encrypt_pdf(
    input_path: str,
    output_path: str,
    user_password: str,
    owner_password: Optional[str] = None,
    progress_callback=None,
    input_password: Optional[str] = None,
) -> bool:
    """Kullanıcı şifresi ile PDF şifreler (açmak için parola gerekir)."""
    try:
        import pikepdf
    except ImportError as e:
        raise Exception("PDF şifreleme için 'pikepdf' gerekli.") from e

    if not user_password:
        raise ValueError("Şifre boş olamaz.")

    try:
        if progress_callback:
            progress_callback(0, 2, "Şifre uygulanıyor...")
        owner = owner_password if owner_password else user_password
        if is_pdf_encrypted(input_path) and not input_password:
            raise Exception(f"PDF şifreleme için kaynak dosya şifresi gerekli: {os.path.basename(input_path)}")
        open_password = (input_password or "").strip()
        target_path = output_path
        temp_output_path = None
        if os.path.abspath(input_path) == os.path.abspath(output_path):
            fd, temp_output_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(output_path) or None)
            os.close(fd)
            target_path = temp_output_path
        with pikepdf.open(input_path, password=open_password) as pdf:
            pdf.save(
                target_path,
                encryption=pikepdf.Encryption(
                    user=user_password,
                    owner=owner,
                    R=6,
                    allow=pikepdf.Permissions(extract=False),
                ),
            )
        if temp_output_path:
            os.replace(temp_output_path, output_path)
        if progress_callback:
            progress_callback(2, 2, "Tamamlandı")
        return True
    except Exception as e:
        raise Exception(f"PDF Şifreleme Hatası: {e}")
    finally:
        if "temp_output_path" in locals() and temp_output_path and os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except Exception:
                pass
